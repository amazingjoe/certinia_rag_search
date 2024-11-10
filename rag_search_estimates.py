from MilvusUtilities import search_milvus, get_milvus_client, get_milvus_collection
from llama_index.embeddings.ibm import WatsonxEmbeddings
from llama_index.core import Settings
from docx import Document
from ibmcos_utilitites import retrieve_file_from_cos
from llm_retriever import query_llm
import os
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

embed_model = WatsonxEmbeddings(model_id=os.getenv("EMBED_MODEL_ID"),
                    url=os.getenv("EMBED_URL"),
                    apikey=os.getenv("EMBED_APIKEY"),
                    project_id=os.getenv("EMBED_PROJECT_ID"))

Settings.embed_model = embed_model

milvus_client = get_milvus_client()
collection = get_milvus_collection(milvus_client,"CertiniaEstimatesRAG")

# Load the collection into memory
collection.load()

def read_full_text_from_file(bucketname,document_name):
    """
    Read the full text from a file given its path.

    Parameters:
    - file_path (str): The path to the file.

    Returns:
    - full_text (str): The full text read from the file.
    """
    if bucketname and document_name:
        file_content = retrieve_file_from_cos(bucketname, document_name)
        if document_name.endswith('.md'):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1')
        elif document_name.endswith('.docx'):
            # Load the file content into a Document object
            doc = Document(BytesIO(file_content))

            # Extract and return the text from the document
            return '\n'.join([para.text for para in doc.paragraphs])
            
            #doc = Document(file_content)
            #return '\n'.join([para.text for para in doc.paragraphs])
    return "Unsupported file format"

def rag_search_and_answer(question):
    context, document_name, bucket = rag_search(question)
    
    import json
    response = {
        "answer": query_llm(context, question),
        "document_name": document_name
    }
    return json.dumps(response)

def rag_search(query):
    embedding = embed_model.get_text_embedding(query)
    results = search_milvus(milvus_client,"CertiniaEstimatesRAG",embedding)

    hit_dict = results[0][0].to_dict()

    bucket = hit_dict['entity']['bucket']
    document_name = hit_dict['entity']['document_name']

    full_text = read_full_text_from_file(bucket,document_name)
    return full_text, document_name, bucket
